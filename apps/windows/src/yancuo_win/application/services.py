"""应用服务：科目、标签、错题、导入、备份、导出。UI 只依赖本层。"""

from __future__ import annotations

import csv
import hashlib
import json
import os
import random
import shutil
import stat
import tempfile
import zipfile
from email.utils import parsedate_to_datetime
from time import monotonic
from urllib.error import HTTPError, URLError
from urllib.request import Request
from dataclasses import dataclass
from datetime import date, datetime, timedelta, timezone
from pathlib import Path
from typing import Any, Iterable

from sqlalchemy import Select, delete, func, or_, select
from sqlalchemy.exc import SQLAlchemyError
from sqlalchemy.orm import Session, selectinload

from yancuo_win.application.bootstrap import RuntimeContext
from yancuo_win.application.sqlite_snapshot import create_sqlite_snapshot
from yancuo_win.assets.object_store import ObjectStore
from yancuo_win.infrastructure.safe_http import safe_urlopen
from yancuo_win.data.ids import new_id
from yancuo_win.data.models import (
    AiJob,
    AiJobItem,
    Asset,
    Chapter,
    ChapterAlias,
    IntakeAsset,
    IntakeCandidateRecord,
    NoteAsset,
    NoteIntakeAsset,
    NoteDocument,
    NoteStudyRecord,
    Problem,
    ProblemOrigin,
    Prompt,
    ReviewItem,
    ReviewPlan,
    ReviewPlanItem,
    ReviewSession,
    ReviewWaitingItem,
    StudyRecord,
    StudySession,
    Subject,
    Tag,
    Version,
    utcnow,
)
from yancuo_win.domain.rules import (
    DomainError,
    assert_transition,
    validate_priority,
    validate_status,
)
from yancuo_win.domain.review_rules import (
    REVIEW_GRADES,
    compute_next_review_at,
    is_due,
    mastery_from_grade,
    validate_grade,
)
from yancuo_win.domain.similarity import normalize_text, text_similarity
from yancuo_win.infrastructure.archive import (
    ArchiveSecurityError,
    iter_regular_files,
    read_regular_file_limited,
    read_zip_member_limited,
    safe_extract_zip,
    validate_zip_members,
)
from yancuo_win.domain.identity import read_identity

MAX_BACKUP_METADATA_BYTES = 4 * 1024 * 1024
MAX_CHAPTER_TEMPLATE_BYTES = 4 * 1024 * 1024
MAX_CHAPTER_TEMPLATE_ITEMS = 10_000


def _sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _verify_local_backup_checksums(root: Path, manifest: dict[str, Any]) -> None:
    if int(manifest.get("version") or 0) < 2:
        return
    declared = manifest.get("checksums")
    if not isinstance(declared, dict):
        raise DomainError("备份 manifest 缺少 checksums 对象")
    actual_files = {
        path.relative_to(root).as_posix(): path
        for path in iter_regular_files(root)
        if path.relative_to(root).as_posix() != "manifest.json"
    }
    if set(declared) != set(actual_files):
        raise DomainError("备份 checksums 与归档载荷不一致")
    for relative_path, path in actual_files.items():
        expected = declared.get(relative_path)
        if (
            not isinstance(expected, str)
            or len(expected) != 64
            or any(character not in "0123456789abcdef" for character in expected)
        ):
            raise DomainError(f"备份 checksum 无效：{relative_path}")
        if _sha256_file(path) != expected:
            raise DomainError(f"备份 checksum 不匹配：{relative_path}")


@dataclass
class ProblemFilter:
    status: str | None = None  # None=非回收站日常；"all"=全部；具体状态；"library"=inbox+active
    subject_id: str | None = None
    chapter_id: str | None = None
    include_descendant_chapters: bool = False
    only_uncategorized: bool = False
    tag_id: str | None = None
    priority: int | None = None
    query: str | None = None
    include_trashed: bool = False
    due_for_review: bool = False
    favorite_only: bool = False
    created_within_days: int | None = None


@dataclass(frozen=True)
class ChapterTreeNode:
    """Stable application-layer projection used by the library knowledge tree."""

    chapter_id: str
    subject_id: str
    parent_id: str | None
    name: str
    sort_order: int
    depth: int
    path_ids: tuple[str, ...]
    path_names: tuple[str, ...]
    direct_problem_count: int
    total_problem_count: int
    children: tuple[ChapterTreeNode, ...]

    @property
    def path_label(self) -> str:
        return " / ".join(self.path_names)


@dataclass(frozen=True)
class CategoryChoice:
    subject_id: str
    subject_name: str
    chapter_id: str | None
    chapter_path: tuple[str, ...] = ()

    @property
    def label(self) -> str:
        suffix = " / ".join(self.chapter_path) if self.chapter_path else "未分类"
        return f"{self.subject_name} / {suffix}"


@dataclass(frozen=True)
class KnowledgeScope:
    key: str
    label: str
    subject_id: str | None = None
    chapter_id: str | None = None
    include_descendants: bool = False
    only_uncategorized: bool = False


class AppServices:
    def __init__(self, runtime: RuntimeContext) -> None:
        self.runtime = runtime
        self.store = ObjectStore(runtime.paths.asset_objects_dir)
        self._review_date_cache: tuple[float, str] | None = None

    def session(self) -> Session:
        return self.runtime.session_factory()

    def _record_sync_change(
        self,
        problem: Problem | str,
        *,
        before: dict[str, Any],
        after: dict[str, Any],
        operation: str = "update",
    ) -> None:
        """将所有正式题目写操作统一登记到增量 Operation 日志。

        同步日志单独开事务，避免把 UI 用例的数据库事务和云端实现耦合在一起。
        """
        from yancuo_win.application.sync_service import SyncService

        SyncService(self.runtime).record_problem_update(
            problem, before=before, after=after, operation=operation
        )

    # ---- catalog ----

    def list_subjects(self) -> list[Subject]:
        with self.session() as s:
            rows = s.scalars(select(Subject).order_by(Subject.sort_order, Subject.name)).all()
            s.expunge_all()
            return list(rows)

    def create_subject(self, name: str, sort_order: int = 0) -> Subject:
        name = name.strip()
        if not name:
            raise DomainError("科目名称不能为空")
        with self.session() as s:
            existing = s.scalar(select(Subject).where(Subject.name == name))
            if existing:
                raise DomainError(f"科目已存在：{name}")
            sub = Subject(id=new_id("sub"), name=name, sort_order=sort_order)
            s.add(sub)
            s.commit()
            s.refresh(sub)
            s.expunge(sub)
            return sub

    def rename_subject(self, subject_id: str, name: str) -> None:
        name = name.strip()
        if not name:
            raise DomainError("科目名称不能为空")
        with self.session() as s:
            sub = s.get(Subject, subject_id)
            if not sub:
                raise DomainError("科目不存在")
            duplicate = s.scalar(
                select(Subject).where(
                    Subject.name == name,
                    Subject.id != subject_id,
                )
            )
            if duplicate:
                raise DomainError(f"科目已存在：{name}")
            sub.name = name
            sub.updated_at = utcnow()
            s.commit()

    def delete_subject(self, subject_id: str) -> None:
        with self.session() as s:
            sub = s.get(Subject, subject_id)
            if not sub:
                return
            chapter_count = s.scalar(
                select(func.count()).select_from(Chapter).where(Chapter.subject_id == subject_id)
            )
            problem_count = s.scalar(
                select(func.count()).select_from(Problem).where(Problem.subject_id == subject_id)
            )
            if chapter_count or problem_count:
                raise DomainError("科目下仍有章节或题目，无法删除")
            s.delete(sub)
            s.commit()

    def reorder_subject(self, subject_id: str, delta: int) -> None:
        with self.session() as s:
            rows = list(
                s.scalars(
                    select(Subject).order_by(
                        Subject.sort_order,
                        Subject.name,
                        Subject.id,
                    )
                ).all()
            )
            index = next((i for i, row in enumerate(rows) if row.id == subject_id), -1)
            if index < 0:
                raise DomainError("科目不存在")
            target = max(0, min(len(rows) - 1, index + delta))
            if target == index:
                return
            row = rows.pop(index)
            rows.insert(target, row)
            for order, subject in enumerate(rows):
                subject.sort_order = order
                subject.updated_at = utcnow()
            s.commit()

    def list_chapters(self, subject_id: str) -> list[Chapter]:
        with self.session() as s:
            rows = s.scalars(
                select(Chapter)
                .where(Chapter.subject_id == subject_id)
                .order_by(Chapter.sort_order, Chapter.name, Chapter.id)
            ).all()
            s.expunge_all()
            return list(rows)

    @staticmethod
    def _chapter_sort_key(chapter: Chapter) -> tuple[int, str, str]:
        return (chapter.sort_order, chapter.name.casefold(), chapter.id)

    def _validate_chapter_parent(
        self,
        session: Session,
        *,
        subject_id: str,
        parent_id: str | None,
        moving_chapter_id: str | None = None,
    ) -> Chapter | None:
        if parent_id is None:
            return None
        parent = session.get(Chapter, parent_id)
        if parent is None:
            raise DomainError("上级章节不存在")
        if parent.subject_id != subject_id:
            raise DomainError("上级章节必须属于同一科目")
        if moving_chapter_id is None:
            return parent
        if parent.id == moving_chapter_id:
            raise DomainError("章节不能成为自己的上级")

        visited: set[str] = set()
        cursor: Chapter | None = parent
        while cursor is not None:
            if cursor.id in visited:
                raise DomainError("现有章节目录包含循环引用，请先修复数据")
            visited.add(cursor.id)
            if cursor.id == moving_chapter_id:
                raise DomainError("不能把章节移动到自己的下级")
            cursor = session.get(Chapter, cursor.parent_id) if cursor.parent_id else None
        return parent

    @staticmethod
    def _ensure_unique_chapter_name(
        session: Session,
        *,
        subject_id: str,
        parent_id: str | None,
        name: str,
        exclude_id: str | None = None,
    ) -> None:
        stmt = select(Chapter).where(
            Chapter.subject_id == subject_id,
            Chapter.parent_id == parent_id,
        )
        for sibling in session.scalars(stmt):
            if sibling.id != exclude_id and sibling.name.casefold() == name.casefold():
                raise DomainError(f"同一层级已存在章节：{name}")

    @staticmethod
    def _normalized_taxonomy_name(name: str) -> str:
        value = name.strip()
        for weak in ("求", "关于", "相关", "基础", "专题"):
            value = value.replace(weak, "")
        return normalize_text(value)

    def _ensure_no_chapter_alias_conflict(
        self, session: Session, *, chapter_id: str | None, name: str, parent_id: str | None = None
    ) -> None:
        normalized = self._normalized_taxonomy_name(name)
        if not normalized:
            return
        for chapter in session.scalars(select(Chapter).where(Chapter.parent_id == parent_id)):
            if (
                chapter.id != chapter_id
                and self._normalized_taxonomy_name(chapter.name) == normalized
            ):
                raise DomainError(f"分类名称与既有章节近义重复：{chapter.name}")
        alias = session.scalar(
            select(ChapterAlias).where(ChapterAlias.normalized_name == normalized)
        )
        if alias is not None and alias.chapter_id != chapter_id:
            raise DomainError("分类名称与既有章节别名近义重复")

    def add_chapter_alias(self, chapter_id: str, name: str) -> ChapterAlias:
        name = name.strip()
        if not name:
            raise DomainError("章节别名不能为空")
        with self.session() as s:
            if s.get(Chapter, chapter_id) is None:
                raise DomainError("章节不存在")
            self._ensure_no_chapter_alias_conflict(s, chapter_id=chapter_id, name=name)
            alias = ChapterAlias(
                id=new_id("chalias"),
                chapter_id=chapter_id,
                name=name,
                normalized_name=self._normalized_taxonomy_name(name),
            )
            s.add(alias)
            s.commit()
            s.refresh(alias)
            s.expunge(alias)
            return alias

    def create_chapter(
        self, subject_id: str, name: str, parent_id: str | None = None, sort_order: int = 0
    ) -> Chapter:
        name = name.strip()
        if not name:
            raise DomainError("章节名称不能为空")
        with self.session() as s:
            if not s.get(Subject, subject_id):
                raise DomainError("科目不存在")
            self._validate_chapter_parent(
                s,
                subject_id=subject_id,
                parent_id=parent_id,
            )
            self._ensure_unique_chapter_name(
                s,
                subject_id=subject_id,
                parent_id=parent_id,
                name=name,
            )
            self._ensure_no_chapter_alias_conflict(
                s, chapter_id=None, name=name, parent_id=parent_id
            )
            ch = Chapter(
                id=new_id("ch"),
                subject_id=subject_id,
                parent_id=parent_id,
                name=name,
                sort_order=sort_order,
            )
            s.add(ch)
            s.commit()
            s.refresh(ch)
            s.expunge(ch)
            return ch

    def rename_chapter(self, chapter_id: str, name: str) -> Chapter:
        name = name.strip()
        if not name:
            raise DomainError("章节名称不能为空")
        with self.session() as s:
            chapter = s.get(Chapter, chapter_id)
            if chapter is None:
                raise DomainError("章节不存在")
            self._ensure_unique_chapter_name(
                s,
                subject_id=chapter.subject_id,
                parent_id=chapter.parent_id,
                name=name,
                exclude_id=chapter.id,
            )
            self._ensure_no_chapter_alias_conflict(
                s, chapter_id=chapter.id, name=name, parent_id=chapter.parent_id
            )
            chapter.name = name
            chapter.updated_at = utcnow()
            s.commit()
            s.refresh(chapter)
            s.expunge(chapter)
            return chapter

    def move_chapter(
        self,
        chapter_id: str,
        parent_id: str | None,
        *,
        sort_order: int | None = None,
    ) -> Chapter:
        with self.session() as s:
            chapter = s.get(Chapter, chapter_id)
            if chapter is None:
                raise DomainError("章节不存在")
            self._validate_chapter_parent(
                s,
                subject_id=chapter.subject_id,
                parent_id=parent_id,
                moving_chapter_id=chapter.id,
            )
            self._ensure_unique_chapter_name(
                s,
                subject_id=chapter.subject_id,
                parent_id=parent_id,
                name=chapter.name,
                exclude_id=chapter.id,
            )
            chapter.parent_id = parent_id
            if sort_order is not None:
                chapter.sort_order = int(sort_order)
            chapter.updated_at = utcnow()
            s.commit()
            s.refresh(chapter)
            s.expunge(chapter)
            return chapter

    def delete_chapter(self, chapter_id: str) -> None:
        with self.session() as s:
            chapter = s.get(Chapter, chapter_id)
            if chapter is None:
                return
            has_children = s.scalar(
                select(func.count()).select_from(Chapter).where(Chapter.parent_id == chapter_id)
            )
            has_problems = s.scalar(
                select(func.count()).select_from(Problem).where(Problem.chapter_id == chapter_id)
            )
            if has_children:
                raise DomainError("章节下仍有子章节，无法删除")
            if has_problems:
                raise DomainError("章节下仍有题目，无法删除")
            s.delete(chapter)
            s.commit()

    def reorder_chapter(self, chapter_id: str, delta: int) -> None:
        with self.session() as s:
            chapter = s.get(Chapter, chapter_id)
            if chapter is None:
                raise DomainError("章节不存在")
            siblings = list(
                s.scalars(
                    select(Chapter).where(
                        Chapter.subject_id == chapter.subject_id,
                        Chapter.parent_id == chapter.parent_id,
                    )
                ).all()
            )
            siblings.sort(key=self._chapter_sort_key)
            index = next(
                (i for i, sibling in enumerate(siblings) if sibling.id == chapter_id),
                -1,
            )
            target = max(0, min(len(siblings) - 1, index + delta))
            if index < 0 or target == index:
                return
            row = siblings.pop(index)
            siblings.insert(target, row)
            for order, sibling in enumerate(siblings):
                sibling.sort_order = order
                sibling.updated_at = utcnow()
            s.commit()

    def list_category_choices(self) -> tuple[CategoryChoice, ...]:
        choices: list[CategoryChoice] = []

        def append_nodes(subject: Subject, nodes: tuple[ChapterTreeNode, ...]) -> None:
            for node in nodes:
                choices.append(
                    CategoryChoice(
                        subject_id=subject.id,
                        subject_name=subject.name,
                        chapter_id=node.chapter_id,
                        chapter_path=node.path_names,
                    )
                )
                append_nodes(subject, node.children)

        for subject in self.list_subjects():
            choices.append(
                CategoryChoice(
                    subject_id=subject.id,
                    subject_name=subject.name,
                    chapter_id=None,
                )
            )
            append_nodes(
                subject,
                self.list_chapter_tree(subject.id, problem_status=None),
            )
        return tuple(choices)

    def list_knowledge_scopes(self) -> tuple[KnowledgeScope, ...]:
        scopes = [KnowledgeScope(key="active", label="全部正式题目")]
        category_choices = self.list_category_choices()
        for subject in self.list_subjects():
            scopes.append(
                KnowledgeScope(
                    key=f"subject:{subject.id}",
                    label=subject.name,
                    subject_id=subject.id,
                )
            )
            scopes.append(
                KnowledgeScope(
                    key=f"uncategorized:{subject.id}",
                    label=f"{subject.name} / 未分类",
                    subject_id=subject.id,
                    only_uncategorized=True,
                )
            )
            for choice in category_choices:
                if choice.subject_id != subject.id or choice.chapter_id is None:
                    continue
                scopes.append(
                    KnowledgeScope(
                        key=f"chapter:{subject.id}:{choice.chapter_id}",
                        label=choice.label,
                        subject_id=subject.id,
                        chapter_id=choice.chapter_id,
                        include_descendants=True,
                    )
                )
        return tuple(scopes)

    @staticmethod
    def filter_for_knowledge_scope(
        scope: KnowledgeScope,
        *,
        query: str | None = None,
    ) -> ProblemFilter:
        return ProblemFilter(
            status="active",
            subject_id=scope.subject_id,
            chapter_id=scope.chapter_id,
            include_descendant_chapters=scope.include_descendants,
            only_uncategorized=scope.only_uncategorized,
            query=query,
        )

    def _chapter_subtree_ids(self, session: Session, chapter_id: str) -> tuple[str, ...]:
        chapter = session.get(Chapter, chapter_id)
        if chapter is None:
            raise DomainError("章节不存在")
        chapters = list(
            session.scalars(select(Chapter).where(Chapter.subject_id == chapter.subject_id)).all()
        )
        children_by_parent: dict[str | None, list[Chapter]] = {}
        for item in chapters:
            children_by_parent.setdefault(item.parent_id, []).append(item)

        result: list[str] = []
        pending = [chapter.id]
        visited: set[str] = set()
        while pending:
            current_id = pending.pop()
            if current_id in visited:
                raise DomainError("章节目录包含循环引用")
            visited.add(current_id)
            result.append(current_id)
            pending.extend(child.id for child in children_by_parent.get(current_id, []))
        return tuple(result)

    def chapter_subtree_ids(self, chapter_id: str) -> tuple[str, ...]:
        """Return a validated chapter subtree for read-side projections."""

        with self.session() as session:
            return self._chapter_subtree_ids(session, chapter_id)

    def list_chapter_tree(
        self,
        subject_id: str,
        *,
        problem_status: str | None = "active",
    ) -> tuple[ChapterTreeNode, ...]:
        """Return a validated recursive chapter projection with aggregate counts."""

        with self.session() as s:
            if s.get(Subject, subject_id) is None:
                raise DomainError("科目不存在")
            chapters = list(
                s.scalars(select(Chapter).where(Chapter.subject_id == subject_id)).all()
            )
            chapter_by_id = {chapter.id: chapter for chapter in chapters}
            children_by_parent: dict[str | None, list[Chapter]] = {}
            for chapter in chapters:
                if chapter.parent_id is not None:
                    parent = chapter_by_id.get(chapter.parent_id)
                    if parent is None or parent.subject_id != subject_id:
                        raise DomainError("章节目录包含无效的上级引用")
                children_by_parent.setdefault(chapter.parent_id, []).append(chapter)
            for children in children_by_parent.values():
                children.sort(key=self._chapter_sort_key)

            count_stmt = (
                select(Problem.chapter_id, func.count())
                .where(
                    Problem.chapter_id.in_(tuple(chapter_by_id)),
                )
                .group_by(Problem.chapter_id)
            )
            if problem_status is not None:
                count_stmt = count_stmt.where(Problem.status == validate_status(problem_status))
            direct_counts = {
                str(chapter_id): int(count)
                for chapter_id, count in s.execute(count_stmt)
                if chapter_id is not None
            }

            visited: set[str] = set()
            active_path: set[str] = set()

            def build(
                chapter: Chapter,
                path_ids: tuple[str, ...],
                path_names: tuple[str, ...],
            ) -> ChapterTreeNode:
                if chapter.id in active_path:
                    raise DomainError("章节目录包含循环引用")
                if chapter.id in visited:
                    raise DomainError("章节目录包含重复引用")
                active_path.add(chapter.id)
                current_ids = (*path_ids, chapter.id)
                current_names = (*path_names, chapter.name)
                children = tuple(
                    build(child, current_ids, current_names)
                    for child in children_by_parent.get(chapter.id, [])
                )
                active_path.remove(chapter.id)
                visited.add(chapter.id)
                direct_count = direct_counts.get(chapter.id, 0)
                return ChapterTreeNode(
                    chapter_id=chapter.id,
                    subject_id=chapter.subject_id,
                    parent_id=chapter.parent_id,
                    name=chapter.name,
                    sort_order=chapter.sort_order,
                    depth=len(path_ids),
                    path_ids=current_ids,
                    path_names=current_names,
                    direct_problem_count=direct_count,
                    total_problem_count=direct_count
                    + sum(child.total_problem_count for child in children),
                    children=children,
                )

            roots = tuple(build(root, (), ()) for root in children_by_parent.get(None, []))
            if len(visited) != len(chapters):
                raise DomainError("章节目录包含无法从根节点访问的循环引用")
            return roots

    def export_chapter_template(self, subject_id: str, dest: Path) -> Path:
        with self.session() as s:
            sub = s.get(Subject, subject_id)
            if not sub:
                raise DomainError("科目不存在")
            subject_name = sub.name

        chapters_payload: list[dict[str, Any]] = []

        def append_nodes(nodes: tuple[ChapterTreeNode, ...]) -> None:
            for node in nodes:
                chapters_payload.append(
                    {
                        "name": node.name,
                        "parent_path": list(node.path_names[:-1]),
                        "sort_order": node.sort_order,
                    }
                )
                append_nodes(node.children)

        append_nodes(self.list_chapter_tree(subject_id, problem_status=None))
        payload = {
            "format": "yancuo-chapter-template",
            "version": 2,
            "subject": {"name": subject_name},
            "chapters": chapters_payload,
        }
        dest.parent.mkdir(parents=True, exist_ok=True)
        dest.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        return dest

    def import_chapter_template(self, path: Path) -> str:
        try:
            raw = json.loads(
                read_regular_file_limited(path, max_bytes=MAX_CHAPTER_TEMPLATE_BYTES).decode(
                    "utf-8"
                )
            )
        except (ArchiveSecurityError, UnicodeDecodeError, json.JSONDecodeError) as exc:
            raise DomainError(f"章节模板无法读取或解析：{exc}") from exc
        if not isinstance(raw, dict):
            raise DomainError("章节模板根节点必须是对象")
        if raw.get("format") != "yancuo-chapter-template":
            raise DomainError("不是有效的章节模板")
        version = int(raw.get("version") or 1)
        if version not in {1, 2}:
            raise DomainError(f"不支持的章节模板版本：{version}")
        chapters = raw.get("chapters")
        if not isinstance(chapters, list) or len(chapters) > MAX_CHAPTER_TEMPLATE_ITEMS:
            raise DomainError("章节模板 chapters 必须是未超限数组")
        subject = raw.get("subject")
        if not isinstance(subject, dict) or not str(subject.get("name") or "").strip():
            raise DomainError("章节模板缺少科目名称")
        subject_name = str(subject["name"])
        with self.session() as s:
            sub = s.scalar(select(Subject).where(Subject.name == subject_name))
            if not sub:
                sub = Subject(id=new_id("sub"), name=subject_name)
                s.add(sub)
                s.flush()
            existing = list(s.scalars(select(Chapter).where(Chapter.subject_id == sub.id)).all())
            if version == 1:
                name_to_id = {chapter.name: chapter.id for chapter in existing}
                for item in chapters:
                    name = str(item["name"]).strip()
                    if not name or name in name_to_id:
                        continue
                    parent_name = item.get("parent_name")
                    parent_id = name_to_id.get(parent_name) if parent_name else None
                    ch = Chapter(
                        id=new_id("ch"),
                        subject_id=sub.id,
                        parent_id=parent_id,
                        name=name,
                        sort_order=int(item.get("sort_order") or 0),
                    )
                    s.add(ch)
                    s.flush()
                    name_to_id[name] = ch.id
            else:
                chapter_by_id = {chapter.id: chapter for chapter in existing}
                path_cache: dict[str, tuple[str, ...]] = {}
                active: set[str] = set()

                def resolve_path(chapter: Chapter) -> tuple[str, ...]:
                    cached = path_cache.get(chapter.id)
                    if cached is not None:
                        return cached
                    if chapter.id in active:
                        raise DomainError("现有章节目录包含循环引用")
                    active.add(chapter.id)
                    if chapter.parent_id is None:
                        result = (chapter.name,)
                    else:
                        parent = chapter_by_id.get(chapter.parent_id)
                        if parent is None:
                            raise DomainError("现有章节目录包含无效的上级引用")
                        result = (*resolve_path(parent), chapter.name)
                    active.remove(chapter.id)
                    path_cache[chapter.id] = result
                    return result

                path_to_id: dict[tuple[str, ...], str] = {}
                for chapter in existing:
                    chapter_path = resolve_path(chapter)
                    if chapter_path in path_to_id:
                        raise DomainError("现有章节目录包含重复知识路径")
                    path_to_id[chapter_path] = chapter.id

                items = sorted(
                    chapters,
                    key=lambda item: len(item.get("parent_path") or []),
                )
                for item in items:
                    name = str(item["name"]).strip()
                    if not name:
                        raise DomainError("章节模板包含空名称")
                    raw_parent_path = item.get("parent_path") or []
                    if not isinstance(raw_parent_path, list):
                        raise DomainError("章节模板 parent_path 必须是数组")
                    parent_path = tuple(str(part).strip() for part in raw_parent_path)
                    if any(not part for part in parent_path):
                        raise DomainError("章节模板包含空的上级路径")
                    chapter_path = (*parent_path, name)
                    if chapter_path in path_to_id:
                        continue
                    parent_id = path_to_id.get(parent_path) if parent_path else None
                    if parent_path and parent_id is None:
                        raise DomainError(f"章节模板缺少上级路径：{' / '.join(parent_path)}")
                    self._ensure_unique_chapter_name(
                        s,
                        subject_id=sub.id,
                        parent_id=parent_id,
                        name=name,
                    )
                    chapter = Chapter(
                        id=new_id("ch"),
                        subject_id=sub.id,
                        parent_id=parent_id,
                        name=name,
                        sort_order=int(item.get("sort_order") or 0),
                    )
                    s.add(chapter)
                    s.flush()
                    path_to_id[chapter_path] = chapter.id
            s.commit()
            return sub.id

    # ---- tags ----

    def list_tags(self) -> list[Tag]:
        with self.session() as s:
            rows = s.scalars(select(Tag).order_by(Tag.name)).all()
            s.expunge_all()
            return list(rows)

    def create_tag(self, name: str, color: str | None = None) -> Tag:
        name = name.strip()
        if not name:
            raise DomainError("标签名称不能为空")
        with self.session() as s:
            existing = s.scalar(select(Tag).where(Tag.name == name))
            if existing:
                raise DomainError(f"标签已存在：{name}")
            tag = Tag(id=new_id("tag"), name=name, color=color, is_system=False)
            s.add(tag)
            s.commit()
            s.refresh(tag)
            s.expunge(tag)
            return tag

    def delete_tag(self, tag_id: str) -> None:
        with self.session() as s:
            tag = s.get(Tag, tag_id)
            if not tag:
                return
            if tag.is_system:
                raise DomainError("系统标签不可删除")
            s.delete(tag)
            s.commit()

    def set_problem_tags(self, problem_id: str, tag_ids: list[str]) -> None:
        from yancuo_win.application.sync_service import sync_snapshot

        with self.session() as s:
            problem = s.scalars(
                select(Problem).where(Problem.id == problem_id).options(selectinload(Problem.tags))
            ).first()
            if not problem:
                raise DomainError("题目不存在")
            before = sync_snapshot(problem)
            tags = list(s.scalars(select(Tag).where(Tag.id.in_(tag_ids))).all()) if tag_ids else []
            problem.tags = tags
            problem.updated_at = utcnow()
            self._add_version(s, problem, source="manual", summary="更新标签")
            after = sync_snapshot(problem, [t.name for t in tags])
            s.commit()
            s.refresh(problem)
            s.expunge(problem)
        self._record_sync_change(problem, before=before, after=after)

    # ---- problems ----

    def _problem_query(
        self,
        filt: ProblemFilter,
        *,
        chapter_ids: tuple[str, ...] | None = None,
    ) -> Select[tuple[Problem]]:
        stmt = select(Problem).options(
            selectinload(Problem.tags),
            selectinload(Problem.assets),
        )
        if filt.status == "library":
            stmt = stmt.where(Problem.status.in_(("inbox", "active")))
        elif filt.status == "all":
            pass
        elif filt.status:
            stmt = stmt.where(Problem.status == validate_status(filt.status))
        elif not filt.include_trashed:
            stmt = stmt.where(Problem.status != "trashed")

        if filt.subject_id:
            stmt = stmt.where(Problem.subject_id == filt.subject_id)
        if filt.only_uncategorized:
            stmt = stmt.where(Problem.chapter_id.is_(None))
        elif chapter_ids is not None:
            stmt = stmt.where(Problem.chapter_id.in_(chapter_ids))
        elif filt.chapter_id:
            stmt = stmt.where(Problem.chapter_id == filt.chapter_id)
        if filt.priority is not None:
            stmt = stmt.where(Problem.priority == filt.priority)
        if filt.tag_id:
            stmt = stmt.where(Problem.tags.any(Tag.id == filt.tag_id))
        if filt.query:
            q = f"%{filt.query.strip()}%"
            stmt = stmt.where(
                or_(
                    Problem.title.ilike(q),
                    Problem.question_markdown.ilike(q),
                    Problem.correct_answer.ilike(q),
                    Problem.notes.ilike(q),
                    Problem.source_book.ilike(q),
                    Problem.original_number.ilike(q),
                )
            )
        if filt.due_for_review:
            # 正式题库中到期或从未安排复习的题
            stmt = stmt.where(Problem.status == "active")
        if filt.favorite_only:
            stmt = stmt.where(Problem.is_favorite.is_(True))
        if filt.created_within_days is not None:
            if filt.created_within_days < 1:
                raise DomainError("最近入库天数必须大于 0")
            stmt = stmt.where(
                Problem.created_at
                >= datetime.now(timezone.utc) - timedelta(days=filt.created_within_days)
            )
        return stmt.order_by(Problem.updated_at.desc())

    def list_problems(self, filt: ProblemFilter | None = None) -> list[Problem]:
        filt = filt or ProblemFilter(status="library")
        with self.session() as s:
            chapter_ids = None
            if filt.chapter_id and filt.include_descendant_chapters:
                chapter_ids = self._chapter_subtree_ids(s, filt.chapter_id)
            rows = list(s.scalars(self._problem_query(filt, chapter_ids=chapter_ids)).all())
            if filt.due_for_review:
                rows = [p for p in rows if p.review_enabled and is_due(p.next_review_at)]
            s.expunge_all()
            return list(rows)

    def list_problems_by_ids(self, problem_ids: Iterable[str]) -> list[Problem]:
        """Load a ranked problem ID result set without losing its order."""

        ordered_ids = list(dict.fromkeys(problem_ids))
        if not ordered_ids:
            return []
        with self.session() as session:
            rows = list(
                session.scalars(
                    select(Problem)
                    .where(Problem.id.in_(ordered_ids))
                    .options(
                        selectinload(Problem.tags),
                        selectinload(Problem.assets),
                    )
                ).all()
            )
            by_id = {problem.id: problem for problem in rows}
            ordered = [by_id[problem_id] for problem_id in ordered_ids if problem_id in by_id]
            session.expunge_all()
            return ordered

    def get_problem(self, problem_id: str) -> Problem | None:
        with self.session() as s:
            problem = s.scalars(
                select(Problem)
                .where(Problem.id == problem_id)
                .options(selectinload(Problem.tags), selectinload(Problem.assets))
            ).first()
            if problem:
                s.expunge_all()
            return problem

    def count_problems(self, status: str | None = None) -> int:
        with self.session() as s:
            stmt = select(func.count()).select_from(Problem)
            if status:
                stmt = stmt.where(Problem.status == status)
            return int(s.scalar(stmt) or 0)

    def create_problem(
        self,
        *,
        title: str | None = None,
        question_markdown: str = "",
        status: str = "inbox",
        subject_id: str | None = None,
        chapter_id: str | None = None,
        priority: int = 3,
    ) -> Problem:
        from yancuo_win.application.sync_service import sync_snapshot

        validate_status(status)
        validate_priority(priority)
        with self.session() as s:
            problem = Problem(
                id=new_id("problem"),
                status=status,
                title=title,
                question_markdown=question_markdown,
                subject_id=subject_id,
                chapter_id=chapter_id,
                priority=priority,
                revision=1,
            )
            s.add(problem)
            s.flush()
            self._add_version(s, problem, source="manual", summary="创建题目", bump=False)
            # 创建操作需要完整快照，供另一台设备首次拉取时建立同一实体。
            after = sync_snapshot(problem, [])
            s.commit()
            s.refresh(problem)
            s.expunge(problem)
        self._record_sync_change(problem, before={}, after=after, operation="create")
        return problem

    def update_problem(
        self, problem_id: str, fields: dict[str, Any], *, summary: str = "编辑题目"
    ) -> Problem:
        from yancuo_win.application.sync_service import sync_snapshot

        with self.session() as s:
            problem = s.scalar(
                select(Problem).where(Problem.id == problem_id).options(selectinload(Problem.tags))
            )
            if not problem:
                raise DomainError("题目不存在")
            before = sync_snapshot(problem)
            allowed = {
                "title",
                "question_markdown",
                "question_latex",
                "question_content_json",
                "user_answer",
                "correct_answer",
                "solution_markdown",
                "error_analysis",
                "notes",
                "source_book",
                "source_year",
                "page_number",
                "original_number",
                "problem_type",
                "subject_id",
                "chapter_id",
                "priority",
                "difficulty",
                "mastery",
                "is_favorite",
                "needs_redo",
                "allow_print",
                "human_confirmed",
            }
            changed = False
            for key, value in fields.items():
                if key not in allowed:
                    continue
                if key == "priority" and value is not None:
                    value = validate_priority(int(value))
                if getattr(problem, key) != value:
                    setattr(problem, key, value)
                    changed = True
            if changed:
                problem.updated_at = utcnow()
                self._add_version(s, problem, source="manual", summary=summary)
            after = sync_snapshot(problem)
            s.commit()
            s.refresh(problem)
            s.expunge(problem)
        if changed:
            self._record_sync_change(problem, before=before, after=after)
        return problem

    def set_problem_status(self, problem_id: str, status: str) -> None:
        from yancuo_win.application.sync_service import sync_snapshot

        with self.session() as s:
            problem = s.scalars(
                select(Problem).where(Problem.id == problem_id).options(selectinload(Problem.tags))
            ).first()
            if not problem:
                raise DomainError("题目不存在")
            before_status = problem.status
            before = sync_snapshot(problem)
            assert_transition(problem.status, status)
            problem.status = status
            problem.updated_at = utcnow()
            if status == "trashed":
                problem.deleted_at = utcnow()
            elif problem.deleted_at is not None:
                problem.deleted_at = None
            self._add_version(s, problem, source="manual", summary=f"状态 → {status}")
            after = sync_snapshot(problem)
            s.commit()
            s.refresh(problem)
            s.expunge(problem)
        operation = "update"
        if status == "trashed":
            operation = "delete"
        elif before_status == "trashed":
            operation = "undelete"
        self._record_sync_change(problem, before=before, after=after, operation=operation)
        if status == "trashed":
            self.remove_review_references("problem", [problem_id])

    def trash_problem(self, problem_id: str) -> None:
        self.set_problem_status(problem_id, "trashed")

    def restore_problem(self, problem_id: str, to_status: str = "inbox") -> None:
        if to_status not in {"inbox", "active"}:
            raise DomainError("恢复目标只能是 inbox 或 active")
        self.set_problem_status(problem_id, to_status)

    def purge_trashed(self) -> int:
        """Permanently delete trashed problems and their dependent workflow data.

        AI and review rows reference problems/assets without database-level cascade
        rules in schema v4.  Remove those rows first so the whole purge remains one
        atomic transaction.  Object-store files are removed only after commit and
        only when no surviving Asset row still references the same relative path.
        """

        relative_paths: set[str] = set()
        try:
            with self.session() as s:
                rows = list(
                    s.scalars(
                        select(Problem)
                        .where(Problem.status == "trashed")
                        .options(
                            selectinload(Problem.tags),
                            selectinload(Problem.assets),
                            selectinload(Problem.versions),
                        )
                    ).all()
                )
                if not rows:
                    return 0

                problem_ids = [problem.id for problem in rows]
                assets = [asset for problem in rows for asset in problem.assets]
                asset_ids = [asset.id for asset in assets]
                relative_paths = {asset.relative_path for asset in assets}

                item_scope = AiJobItem.problem_id.in_(problem_ids)
                if asset_ids:
                    item_scope = or_(item_scope, AiJobItem.asset_id.in_(asset_ids))
                affected_job_ids = set(s.scalars(select(AiJobItem.job_id).where(item_scope)).all())
                affected_session_ids = set(
                    s.scalars(
                        select(ReviewItem.session_id).where(ReviewItem.problem_id.in_(problem_ids))
                    ).all()
                )

                s.execute(delete(ReviewItem).where(ReviewItem.problem_id.in_(problem_ids)))
                s.execute(delete(AiJobItem).where(item_scope))
                s.execute(delete(ProblemOrigin).where(ProblemOrigin.problem_id.in_(problem_ids)))
                for candidate in s.scalars(
                    select(IntakeCandidateRecord).where(
                        IntakeCandidateRecord.problem_id.in_(problem_ids)
                    )
                ).all():
                    candidate.problem_id = None
                s.flush()

                for problem in rows:
                    problem.tags.clear()
                    s.delete(problem)
                s.flush()

                for session_id in affected_session_ids:
                    review_session = s.get(ReviewSession, session_id)
                    if review_session is None:
                        continue
                    has_items = s.scalar(
                        select(func.count(ReviewItem.id)).where(ReviewItem.session_id == session_id)
                    )
                    if not has_items:
                        s.delete(review_session)

                for job_id in affected_job_ids:
                    job = s.get(AiJob, job_id)
                    if job is None:
                        continue
                    remaining = list(
                        s.scalars(select(AiJobItem).where(AiJobItem.job_id == job_id)).all()
                    )
                    if remaining:
                        job.total_items = len(remaining)
                        job.done_items = sum(item.status == "done" for item in remaining)
                        job.failed_items = sum(item.status == "failed" for item in remaining)
                        continue

                    surviving_review_items = False
                    for review_session in s.scalars(
                        select(ReviewSession).where(ReviewSession.job_id == job_id)
                    ).all():
                        has_items = s.scalar(
                            select(func.count(ReviewItem.id)).where(
                                ReviewItem.session_id == review_session.id
                            )
                        )
                        if has_items:
                            surviving_review_items = True
                        else:
                            s.delete(review_session)

                    if surviving_review_items:
                        # One image can yield several candidate problems.  The
                        # original job item may belong to a rejected/trashed
                        # candidate while sibling review items are still open.
                        # Keep the job as their stable intake-session anchor.
                        job.total_items = 0
                        job.done_items = 0
                        job.failed_items = 0
                        continue

                    # review_sessions.job_id has no ON DELETE cascade in schema v4;
                    # persist detach/delete decisions before removing the job.
                    s.flush()
                    prompt_key = job.prompt_key
                    s.delete(job)
                    s.flush()
                    if prompt_key == f"intake_{job_id}":
                        prompt_in_use = s.scalar(
                            select(func.count(AiJob.id)).where(AiJob.prompt_key == prompt_key)
                        )
                        if not prompt_in_use:
                            prompt = s.scalar(select(Prompt).where(Prompt.key == prompt_key))
                            if prompt is not None:
                                s.delete(prompt)

                count = len(rows)
                s.commit()
        except SQLAlchemyError as exc:
            raise DomainError("清空回收站失败，所有删除操作均已回滚") from exc

        self._remove_unreferenced_asset_files(relative_paths)
        return count

    def _remove_unreferenced_asset_files(self, relative_paths: set[str]) -> None:
        """Best-effort cleanup after database references have been committed."""

        objects_root = self.store.objects_root.resolve()
        with self.session() as s:
            referenced = {
                path
                for path in relative_paths
                if s.scalar(select(func.count(Asset.id)).where(Asset.relative_path == path))
                or s.scalar(
                    select(func.count(IntakeAsset.id)).where(IntakeAsset.relative_path == path)
                )
                or s.scalar(select(func.count(NoteAsset.id)).where(NoteAsset.relative_path == path))
                or s.scalar(
                    select(func.count(NoteIntakeAsset.id)).where(
                        NoteIntakeAsset.relative_path == path
                    )
                )
            }

        for relative_path in relative_paths - referenced:
            path = self.store.resolve(relative_path)
            try:
                path.relative_to(objects_root)
            except ValueError:
                self.runtime.logger.warning("skip unsafe asset cleanup path: %s", relative_path)
                continue
            try:
                if path.is_file():
                    path.chmod(path.stat().st_mode | stat.S_IWRITE)
                    path.unlink()
                if path.parent != objects_root:
                    path.parent.rmdir()
            except OSError as exc:
                self.runtime.logger.warning("orphan asset cleanup failed for %s: %s", path, exc)

    def promote_to_active(self, problem_id: str) -> None:
        self.set_problem_status(problem_id, "active")

    # ---- review ----

    def list_due_reviews(self) -> list[Problem]:
        return self.list_problems(ProblemFilter(status="active", due_for_review=True))

    def review_plan_date(self) -> str:
        """Use a network-confirmed UTC date when available, else local Shanghai time."""

        now = monotonic()
        if self._review_date_cache is not None:
            expires_at, cached_date = self._review_date_cache
            if now < expires_at:
                return cached_date
        result: str | None = None
        try:
            request = Request("https://www.cloudflare.com", method="HEAD")
            with safe_urlopen(request, timeout=2) as response:
                header = response.headers.get("Date")
            if header:
                result = (
                    parsedate_to_datetime(header)
                    .astimezone(timezone(timedelta(hours=8)))
                    .date()
                    .isoformat()
                )
        except (HTTPError, URLError, OSError, TypeError, ValueError):
            pass
        if result is None:
            result = datetime.now(timezone(timedelta(hours=8))).date().isoformat()
        self._review_date_cache = (now + 300, result)
        return result

    @staticmethod
    def _validate_review_content_type(content_type: str) -> str:
        if content_type not in {"problem", "note"}:
            raise DomainError("复习内容类型必须是题目或笔记")
        return content_type

    def add_to_review_waiting_queue(self, content_type: str, source_ids: Iterable[str]) -> int:
        content_type = self._validate_review_content_type(content_type)
        ids = list(dict.fromkeys(source_ids))
        with self.session() as session:
            existing = (
                set(
                    session.scalars(
                        select(ReviewWaitingItem.source_id).where(
                            ReviewWaitingItem.content_type == content_type,
                            ReviewWaitingItem.source_id.in_(ids),
                        )
                    ).all()
                )
                if ids
                else set()
            )
            for source_id in ids:
                if source_id not in existing:
                    session.add(
                        ReviewWaitingItem(
                            id=new_id("review_waiting"),
                            content_type=content_type,
                            source_id=source_id,
                        )
                    )
            session.commit()
        return len(ids) - len(existing)

    def remove_from_review_waiting_queue(self, content_type: str, source_ids: Iterable[str]) -> int:
        content_type = self._validate_review_content_type(content_type)
        ids = list(dict.fromkeys(source_ids))
        with self.session() as session:
            result = session.execute(
                delete(ReviewWaitingItem).where(
                    ReviewWaitingItem.content_type == content_type,
                    ReviewWaitingItem.source_id.in_(ids),
                )
            )
            session.commit()
            return int(result.rowcount or 0)

    def clear_review_waiting_queue(self, content_type: str) -> int:
        content_type = self._validate_review_content_type(content_type)
        with self.session() as session:
            result = session.execute(
                delete(ReviewWaitingItem).where(ReviewWaitingItem.content_type == content_type)
            )
            session.commit()
            return int(result.rowcount or 0)

    def list_review_waiting_ids(self, content_type: str) -> list[str]:
        content_type = self._validate_review_content_type(content_type)
        with self.session() as session:
            return list(
                session.scalars(
                    select(ReviewWaitingItem.source_id)
                    .where(ReviewWaitingItem.content_type == content_type)
                    .order_by(ReviewWaitingItem.created_at)
                ).all()
            )

    def create_review_plan_from_waiting_queue(
        self,
        content_type: str,
        name: str,
        source_ids: Iterable[str] | None = None,
    ) -> ReviewPlan:
        content_type = self._validate_review_content_type(content_type)
        name = name.strip()
        if not name:
            raise DomainError("请为复习计划命名")
        with self.session() as session:
            waiting = list(
                session.scalars(
                    select(ReviewWaitingItem)
                    .where(ReviewWaitingItem.content_type == content_type)
                    .order_by(ReviewWaitingItem.created_at)
                ).all()
            )
            if not waiting:
                raise DomainError("等待队列为空")
            if source_ids is not None:
                requested_order = list(dict.fromkeys(source_ids))
                waiting_by_source = {item.source_id: item for item in waiting}
                if set(requested_order) != set(waiting_by_source):
                    raise DomainError("计划草稿已变化，请刷新后再创建")
                waiting = [waiting_by_source[source_id] for source_id in requested_order]
            plan = ReviewPlan(
                id=new_id("review_plan"), name=name, content_type=content_type, kind="explicit"
            )
            session.add(plan)
            session.flush()
            for order, item in enumerate(waiting):
                session.add(
                    ReviewPlanItem(
                        id=new_id("review_plan_item"),
                        plan_id=plan.id,
                        source_id=item.source_id,
                        sort_order=order,
                    )
                )
            session.execute(
                delete(ReviewWaitingItem).where(ReviewWaitingItem.content_type == content_type)
            )
            session.commit()
            session.refresh(plan)
            session.expunge(plan)
            return plan

    def list_review_plans(self, content_type: str | None = None) -> list[ReviewPlan]:
        with self.session() as session:
            statement = (
                select(ReviewPlan)
                .options(selectinload(ReviewPlan.items))
                .order_by(ReviewPlan.updated_at.desc())
            )
            if content_type:
                statement = statement.where(
                    ReviewPlan.content_type == self._validate_review_content_type(content_type)
                )
            rows = list(session.scalars(statement).all())
            session.expunge_all()
            return rows

    def get_review_plan(self, plan_id: str) -> ReviewPlan | None:
        with self.session() as session:
            plan = session.scalar(
                select(ReviewPlan)
                .where(ReviewPlan.id == plan_id)
                .options(selectinload(ReviewPlan.items))
            )
            if plan is not None:
                session.expunge(plan)
            return plan

    def add_to_daily_review_plan(
        self, content_type: str, source_id: str, plan_date: str
    ) -> ReviewPlan:
        content_type = self._validate_review_content_type(content_type)
        with self.session() as session:
            plan = session.scalar(
                select(ReviewPlan).where(
                    ReviewPlan.content_type == content_type,
                    ReviewPlan.kind == "daily",
                    ReviewPlan.plan_date == plan_date,
                )
            )
            if plan is None:
                display_date = date.fromisoformat(plan_date)
                plan = ReviewPlan(
                    id=new_id("review_plan"),
                    name=f"{display_date.year}年{display_date.month}月{display_date.day}日 复习计划",
                    content_type=content_type,
                    kind="daily",
                    plan_date=plan_date,
                )
                session.add(plan)
                session.flush()
            exists = session.scalar(
                select(ReviewPlanItem.id).where(
                    ReviewPlanItem.plan_id == plan.id, ReviewPlanItem.source_id == source_id
                )
            )
            if not exists:
                next_order = int(
                    session.scalar(
                        select(func.count(ReviewPlanItem.id)).where(
                            ReviewPlanItem.plan_id == plan.id
                        )
                    )
                    or 0
                )
                session.add(
                    ReviewPlanItem(
                        id=new_id("review_plan_item"),
                        plan_id=plan.id,
                        source_id=source_id,
                        sort_order=next_order,
                    )
                )
            session.commit()
            session.refresh(plan)
            session.expunge(plan)
            return plan

    def remove_review_references(self, content_type: str, source_ids: Iterable[str]) -> None:
        content_type = self._validate_review_content_type(content_type)
        ids = list(dict.fromkeys(source_ids))
        if not ids:
            return
        with self.session() as session:
            plan_ids = select(ReviewPlan.id).where(ReviewPlan.content_type == content_type)
            session.execute(
                delete(ReviewWaitingItem).where(
                    ReviewWaitingItem.content_type == content_type,
                    ReviewWaitingItem.source_id.in_(ids),
                )
            )
            session.execute(
                delete(ReviewPlanItem).where(
                    ReviewPlanItem.plan_id.in_(plan_ids), ReviewPlanItem.source_id.in_(ids)
                )
            )
            session.commit()

    def prepare_study_queue(
        self,
        *,
        scope: str = "due",
        problem_types: set[str] | None = None,
        order: str = "scheduled",
        limit: int | None = None,
    ) -> list[Problem]:
        """Build a review queue without mutating scheduling state."""

        if scope == "due":
            queue = self.list_due_reviews()
        elif scope == "active":
            queue = [
                problem
                for problem in self.list_problems(ProblemFilter(status="active"))
                if problem.review_enabled
            ]
        elif scope == "unreviewed":
            queue = [
                problem
                for problem in self.list_problems(ProblemFilter(status="active"))
                if problem.review_enabled and not problem.review_count
            ]
        else:
            raise DomainError("不支持的复习范围")

        if problem_types:
            queue = [
                problem
                for problem in queue
                if (problem.problem_type or "未标注").strip() in problem_types
            ]
        if order == "random":
            random.shuffle(queue)
        elif order == "scheduled":
            queue.sort(
                key=lambda problem: (
                    problem.next_review_at is None,
                    problem.next_review_at,
                    problem.created_at,
                )
            )
        else:
            raise DomainError("不支持的复习顺序")
        if limit is not None:
            if limit < 1:
                raise DomainError("复习数量必须大于 0")
            queue = queue[:limit]
        return queue

    def start_study_session(
        self, *, selection: dict[str, Any] | None = None, problem_ids: list[str] | None = None
    ) -> tuple[StudySession, list[Problem]]:
        if problem_ids is not None:
            queue = self.list_problems_by_ids(problem_ids)
        else:
            queue = self.list_due_reviews()
        criteria = selection or {"kind": "due_reviews", "timezone": "Asia/Shanghai"}
        criteria = {**criteria, "problem_ids": [problem.id for problem in queue]}
        study_session = StudySession(
            id=new_id("study"),
            selection_json=json.dumps(criteria, ensure_ascii=False, sort_keys=True),
            problem_count=len(queue),
            status="active",
        )
        with self.session() as s:
            s.add(study_session)
            s.commit()
            s.refresh(study_session)
            s.expunge(study_session)
        return study_session, queue

    def record_review(
        self,
        problem_id: str,
        grade: int,
        *,
        study_session_id: str | None = None,
        answer_viewed_at: datetime | None = None,
        answered_at: datetime | None = None,
    ) -> dict[str, Any]:
        """记录复习结果并安排下次日期。不自动删除任何题目。"""
        from yancuo_win.application.sync_service import sync_snapshot

        grade = validate_grade(grade)
        graded_at = utcnow()
        next_at = compute_next_review_at(grade, from_dt=graded_at)
        interval_days = (next_at - graded_at).days + 1
        with self.session() as s:
            problem = s.scalars(
                select(Problem).where(Problem.id == problem_id).options(selectinload(Problem.tags))
            ).first()
            if not problem:
                raise DomainError("题目不存在")
            if problem.status == "trashed":
                raise DomainError("回收站题目不可复习")
            if not problem.review_enabled:
                raise DomainError("题目已暂停复习，请先恢复")
            if study_session_id:
                study_session = s.get(StudySession, study_session_id)
                if study_session is None or study_session.status != "active":
                    raise DomainError("学习会话不可用")
            before = sync_snapshot(problem)
            problem.mastery = mastery_from_grade(grade)
            problem.next_review_at = next_at
            problem.review_count = int(problem.review_count or 0) + 1
            problem.updated_at = utcnow()
            if problem.status == "inbox":
                # 复习过的题进入正式库更合理
                problem.status = "active"
            s.add(
                StudyRecord(
                    id=new_id("study_record"),
                    study_session_id=study_session_id,
                    problem_id=problem.id,
                    answer_viewed_at=answer_viewed_at,
                    answered_at=answered_at,
                    grade=grade,
                    graded_at=graded_at,
                    interval_days=interval_days,
                    next_review_at=next_at,
                )
            )
            self._add_version(
                s,
                problem,
                source="review",
                summary=f"复习打分 {grade}（{REVIEW_GRADES[grade]}）",
            )
            after = sync_snapshot(problem)
            s.commit()
            s.refresh(problem)
            s.expunge(problem)
        self._record_sync_change(problem, before=before, after=after)
        return {
            "problem_id": problem_id,
            "grade": grade,
            "label": REVIEW_GRADES[grade],
            "next_review_at": next_at.isoformat(),
            "interval_days": interval_days,
            "review_count": problem.review_count,
        }

    def finish_study_session(
        self, study_session_id: str, *, cancelled: bool = False
    ) -> dict[str, Any]:
        with self.session() as s:
            study_session = s.get(StudySession, study_session_id)
            if study_session is None:
                raise DomainError("学习会话不存在")
            if study_session.status == "active":
                study_session.status = "cancelled" if cancelled else "completed"
                study_session.ended_at = utcnow()
                s.commit()
            records = list(
                s.scalars(
                    select(StudyRecord).where(StudyRecord.study_session_id == study_session.id)
                ).all()
            )
            result = {
                "session_id": study_session.id,
                "status": study_session.status,
                "problem_count": study_session.problem_count,
                "completed_count": len(records),
                "remaining_count": max(0, study_session.problem_count - len(records)),
                "grades": {
                    grade: sum(record.grade == grade for record in records)
                    for grade in REVIEW_GRADES
                },
            }
            s.expunge_all()
            return result

    def set_review_enabled(self, problem_id: str, enabled: bool) -> None:
        with self.session() as s:
            problem = s.get(Problem, problem_id)
            if problem is None:
                raise DomainError("题目不存在")
            if problem.status == "trashed":
                raise DomainError("回收站题目不可设置复习")
            problem.review_enabled = enabled
            if enabled and problem.next_review_at is None:
                problem.next_review_at = datetime.now(timezone.utc)
            elif not enabled:
                problem.next_review_at = None
            problem.updated_at = utcnow()
            s.commit()

    def study_session_records(self, study_session_id: str) -> list[StudyRecord]:
        with self.session() as s:
            records = list(
                s.scalars(
                    select(StudyRecord)
                    .where(StudyRecord.study_session_id == study_session_id)
                    .order_by(StudyRecord.graded_at)
                ).all()
            )
            s.expunge_all()
            return records

    def review_plan_study_sessions(self, review_plan_id: str) -> list[StudySession]:
        """Return problem-review sessions started from one named plan."""

        with self.session() as session:
            sessions = list(
                session.scalars(select(StudySession).order_by(StudySession.started_at.desc())).all()
            )
            matching = []
            for study_session in sessions:
                try:
                    selection = json.loads(study_session.selection_json)
                except (TypeError, json.JSONDecodeError):
                    continue
                if selection.get("review_plan_id") == review_plan_id:
                    matching.append(study_session)
            session.expunge_all()
            return matching

    def record_note_review(
        self, note_id: str, *, review_plan_id: str | None = None
    ) -> NoteStudyRecord:
        """Persist one completion without applying problem-style scoring rules."""

        with self.session() as session:
            note = session.get(NoteDocument, note_id)
            if note is None or note.status == "trashed":
                raise DomainError("笔记不存在或已移入回收站")
            if review_plan_id and session.get(ReviewPlan, review_plan_id) is None:
                raise DomainError("复习计划不存在")
            record = NoteStudyRecord(
                id=new_id("note_study_record"),
                note_id=note_id,
                review_plan_id=review_plan_id,
            )
            session.add(record)
            session.commit()
            session.refresh(record)
            session.expunge(record)
            return record

    def note_study_records(self, note_id: str) -> list[NoteStudyRecord]:
        with self.session() as session:
            records = list(
                session.scalars(
                    select(NoteStudyRecord)
                    .where(NoteStudyRecord.note_id == note_id)
                    .order_by(NoteStudyRecord.completed_at.desc())
                ).all()
            )
            session.expunge_all()
            return records

    def review_plan_note_records(self, review_plan_id: str) -> list[NoteStudyRecord]:
        """Return note completions recorded from one named plan."""

        with self.session() as session:
            records = list(
                session.scalars(
                    select(NoteStudyRecord)
                    .where(NoteStudyRecord.review_plan_id == review_plan_id)
                    .order_by(NoteStudyRecord.completed_at.desc())
                ).all()
            )
            session.expunge_all()
            return records

    def export_study_session_csv(self, study_session_id: str, dest: Path) -> Path:
        with dest.open("w", encoding="utf-8-sig", newline="") as handle:
            writer = csv.DictWriter(
                handle,
                fieldnames=[
                    "problem_id",
                    "grade",
                    "answer_viewed_at",
                    "answered_at",
                    "graded_at",
                    "interval_days",
                    "next_review_at",
                ],
            )
            writer.writeheader()
            for record in self.study_session_records(study_session_id):
                writer.writerow(
                    {
                        "problem_id": record.problem_id,
                        "grade": record.grade,
                        "answer_viewed_at": record.answer_viewed_at.isoformat()
                        if record.answer_viewed_at
                        else "",
                        "answered_at": record.answered_at.isoformat() if record.answered_at else "",
                        "graded_at": record.graded_at.isoformat(),
                        "interval_days": record.interval_days,
                        "next_review_at": record.next_review_at.isoformat(),
                    }
                )
        return dest

    def export_study_session_share(self, study_session_id: str, dest: Path) -> Path:
        """Privacy-safe aggregate: no question text, answers, sources, or identifiers."""
        summary = self.finish_study_session(study_session_id)
        summary.pop("session_id", None)
        dest.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
        return dest

    def schedule_initial_review(self, problem_id: str) -> None:
        """将题目加入复习队列（下次=今天）。"""
        from yancuo_win.application.sync_service import sync_snapshot

        with self.session() as s:
            problem = s.scalars(
                select(Problem).where(Problem.id == problem_id).options(selectinload(Problem.tags))
            ).first()
            if not problem:
                raise DomainError("题目不存在")
            before = sync_snapshot(problem)
            problem.next_review_at = datetime.now(timezone.utc).replace(
                hour=0, minute=0, second=0, microsecond=0
            )
            if problem.status == "inbox":
                problem.status = "active"
            problem.updated_at = utcnow()
            self._add_version(
                s,
                problem,
                source="review",
                summary="加入复习队列",
            )
            problem.review_enabled = True
            after = sync_snapshot(problem)
            s.commit()
            s.refresh(problem)
            s.expunge(problem)
        self._record_sync_change(problem, before=before, after=after)

    # ---- duplicates ----

    def find_hash_duplicates(self) -> list[dict[str, Any]]:
        """按原图 sha256 分组，仅提示不删除。"""
        with self.session() as s:
            assets = s.scalars(
                select(Asset).where(Asset.role == "original", Asset.problem_id.is_not(None))
            ).all()
            by_hash: dict[str, list[Asset]] = {}
            for a in assets:
                by_hash.setdefault(a.sha256, []).append(a)
            groups = []
            for sha, items in by_hash.items():
                if len(items) < 2:
                    continue
                groups.append(
                    {
                        "sha256": sha,
                        "problem_ids": [a.problem_id for a in items if a.problem_id],
                        "count": len(items),
                    }
                )
            return groups

    def find_text_similar(
        self, problem_id: str, *, threshold: float = 0.85, limit: int = 20
    ) -> list[dict[str, Any]]:
        """文本相似提示，不自动合并/删除。"""
        with self.session() as s:
            target = s.get(Problem, problem_id)
            if not target:
                raise DomainError("题目不存在")
            others = s.scalars(
                select(Problem).where(
                    Problem.id != problem_id,
                    Problem.status.in_(("inbox", "active", "archived")),
                )
            ).all()
            scored = []
            for p in others:
                score = text_similarity(target.question_markdown or "", p.question_markdown or "")
                if score >= threshold:
                    scored.append(
                        {
                            "problem_id": p.id,
                            "title": p.title,
                            "score": round(score, 4),
                        }
                    )
            scored.sort(key=lambda x: x["score"], reverse=True)
            return scored[:limit]

    def batch_update_problems(
        self,
        problem_ids: list[str],
        *,
        subject_id: str | None = None,
        chapter_id: str | None = None,
        priority: int | None = None,
        add_tag_id: str | None = None,
    ) -> int:
        from yancuo_win.application.sync_service import sync_snapshot

        if not problem_ids:
            return 0
        if priority is not None:
            validate_priority(priority)
        updated = 0
        sync_changes: list[tuple[str, dict[str, Any], dict[str, Any]]] = []
        with self.session() as s:
            tag = s.get(Tag, add_tag_id) if add_tag_id else None
            for pid in problem_ids:
                problem = s.scalars(
                    select(Problem).where(Problem.id == pid).options(selectinload(Problem.tags))
                ).first()
                if not problem or problem.status == "trashed":
                    continue
                before = sync_snapshot(problem)
                changed = False
                if subject_id is not None and problem.subject_id != subject_id:
                    problem.subject_id = subject_id
                    changed = True
                if chapter_id is not None and problem.chapter_id != chapter_id:
                    problem.chapter_id = chapter_id
                    changed = True
                if priority is not None and problem.priority != priority:
                    problem.priority = priority
                    changed = True
                if tag is not None and tag not in problem.tags:
                    problem.tags = list(problem.tags) + [tag]
                    changed = True
                if changed:
                    problem.updated_at = utcnow()
                    self._add_version(s, problem, source="manual", summary="批量更新")
                    after = sync_snapshot(problem)
                    sync_changes.append((problem.id, before, after))
                    updated += 1
            s.commit()
        for problem_id, before, after in sync_changes:
            self._record_sync_change(problem_id, before=before, after=after)
        return updated

    def move_problems_to_category(
        self,
        problem_ids: list[str],
        *,
        subject_id: str | None,
        chapter_id: str | None,
    ) -> int:
        """Move problems to one validated catalog destination, including uncategorized."""

        from yancuo_win.application.sync_service import sync_snapshot

        if not problem_ids:
            return 0
        with self.session() as s:
            if subject_id is not None and s.get(Subject, subject_id) is None:
                raise DomainError("目标科目不存在")
            if chapter_id is not None:
                chapter = s.get(Chapter, chapter_id)
                if chapter is None:
                    raise DomainError("目标章节不存在")
                if subject_id is None or chapter.subject_id != subject_id:
                    raise DomainError("目标章节不属于所选科目")

            changes: list[tuple[str, dict[str, Any], dict[str, Any]]] = []
            for problem in s.scalars(
                select(Problem)
                .where(Problem.id.in_(problem_ids))
                .options(selectinload(Problem.tags))
            ):
                if problem.status == "trashed":
                    continue
                if problem.subject_id == subject_id and problem.chapter_id == chapter_id:
                    continue
                before = sync_snapshot(problem)
                problem.subject_id = subject_id
                problem.chapter_id = chapter_id
                problem.updated_at = utcnow()
                self._add_version(
                    s,
                    problem,
                    source="manual",
                    summary="移动题目分类",
                )
                changes.append((problem.id, before, sync_snapshot(problem)))
            s.commit()
        for problem_id, before, after in changes:
            self._record_sync_change(problem_id, before=before, after=after)
        return len(changes)

    def _add_version(
        self,
        session: Session,
        problem: Problem,
        *,
        source: str,
        summary: str,
        bump: bool = True,
    ) -> None:
        if bump:
            problem.revision += 1
        from yancuo_win.application.sync_service import sync_snapshot

        snap = sync_snapshot(problem)
        session.add(
            Version(
                id=new_id("ver"),
                problem_id=problem.id,
                revision=problem.revision,
                source=source,
                summary=summary,
                snapshot_json=json.dumps(snap, ensure_ascii=False),
                created_by=self.runtime.identity.user_id,
            )
        )

    # ---- image import ----

    def import_images(
        self,
        paths: Iterable[Path],
        *,
        into_status: str = "inbox",
        skip_duplicates: bool | None = None,
    ) -> dict[str, Any]:
        from yancuo_win.application.sync_service import sync_snapshot

        validate_status(into_status)
        skip = (
            self.runtime.settings.import_cfg.skip_duplicates
            if skip_duplicates is None
            else skip_duplicates
        )
        created: list[str] = []
        skipped: list[str] = []
        skipped_existing: list[dict[str, str]] = []
        sync_changes: list[tuple[str, dict[str, Any], dict[str, Any]]] = []
        with self.session() as s:
            for path in paths:
                path = Path(path)
                stored = self.store.store_copy(path, role="original")
                if skip:
                    exists = s.scalar(
                        select(Asset)
                        .join(Problem, Problem.id == Asset.problem_id)
                        .where(
                            Asset.sha256 == stored.sha256,
                            Asset.role == "original",
                            Problem.status != "trashed",
                        )
                    )
                    if exists:
                        skipped.append(str(path))
                        skipped_existing.append(
                            {
                                "path": str(path),
                                "sha256": stored.sha256,
                                "existing_problem_id": exists.problem_id or "",
                                "existing_asset_id": exists.id,
                            }
                        )
                        continue
                problem = Problem(
                    id=new_id("problem"),
                    status=into_status,
                    title=path.stem,
                    question_markdown="",
                    revision=1,
                )
                s.add(problem)
                s.flush()
                asset = Asset(
                    id=new_id("asset"),
                    problem_id=problem.id,
                    role="original",
                    sha256=stored.sha256,
                    relative_path=stored.relative_path,
                    mime_type=stored.mime_type,
                    size_bytes=stored.size_bytes,
                    is_immutable=True,
                )
                s.add(asset)
                self._add_version(
                    s, problem, source="import", summary=f"导入图片 {path.name}", bump=False
                )
                sync_changes.append((problem.id, {}, sync_snapshot(problem, [])))
                created.append(problem.id)
            s.commit()
        for problem_id, before, after in sync_changes:
            self._record_sync_change(problem_id, before=before, after=after, operation="create")
        return {
            "created": created,
            "skipped": skipped,
            "skipped_existing": skipped_existing,
            "duplicate_tip": (
                f"检测到 {len(skipped)} 张重复原图，已跳过且未删除旧题" if skipped else ""
            ),
        }

    def import_folder(self, folder: Path, *, recursive: bool | None = None) -> dict[str, Any]:
        folder = Path(folder)
        if not folder.is_dir():
            raise DomainError(f"不是文件夹：{folder}")
        scan = self.runtime.settings.import_cfg.scan_subfolders if recursive is None else recursive
        exts = {e.lower() for e in self.runtime.settings.import_cfg.supported_extensions}
        files: list[Path] = []
        if scan:
            for p in folder.rglob("*"):
                if p.is_file() and p.suffix.lower() in exts and p.suffix.lower() != ".pdf":
                    files.append(p)
        else:
            for p in folder.iterdir():
                if p.is_file() and p.suffix.lower() in exts and p.suffix.lower() != ".pdf":
                    files.append(p)
        files.sort()
        return self.import_images(files)

    def attach_original_image(self, problem_id: str, path: Path) -> Asset:
        with self.session() as s:
            problem = s.get(Problem, problem_id)
            if not problem:
                raise DomainError("题目不存在")
            stored = self.store.store_copy(path, role="original")
            asset = Asset(
                id=new_id("asset"),
                problem_id=problem.id,
                role="original",
                sha256=stored.sha256,
                relative_path=stored.relative_path,
                mime_type=stored.mime_type,
                size_bytes=stored.size_bytes,
                is_immutable=True,
            )
            s.add(asset)
            problem.updated_at = utcnow()
            self._add_version(s, problem, source="import", summary="附加原图")
            s.commit()
            s.refresh(asset)
            s.expunge(asset)
            return asset

    def try_overwrite_original(self, asset_id: str) -> None:
        """供测试锁定：原图不可覆盖。"""
        with self.session() as s:
            asset = s.get(Asset, asset_id)
            if not asset:
                raise DomainError("资源不存在")
            self.store.assert_can_replace(asset.role, asset.is_immutable)

    # ---- backup ----

    def create_backup(self, dest_zip: Path | None = None) -> Path:
        stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        dest = dest_zip or (self.runtime.paths.backup_dir / f"yancuo-backup-{stamp}.zip")
        dest.parent.mkdir(parents=True, exist_ok=True)
        db_path = self.runtime.paths.database
        asset_dir = self.runtime.paths.asset_dir
        identity = self.runtime.paths.identity_file

        # 释放长连接，再通过 SQLite API 合并已提交的 WAL 页。
        self.runtime.engine.dispose()

        manifest = {
            "format": "yancuo-local-backup",
            "version": 2,
            "created_at": datetime.now(timezone.utc).isoformat(),
            "database_id": self.runtime.identity.database_id,
            "schema_version": self.runtime.schema_version,
        }
        self.runtime.paths.cache_dir.mkdir(parents=True, exist_ok=True)
        staging = Path(tempfile.mkdtemp(prefix="backup-export-", dir=self.runtime.paths.cache_dir))
        archive_temp: Path | None = None
        try:
            snapshot = staging / "error_book.db"
            create_sqlite_snapshot(db_path, snapshot)
            backup_identity_payload: str | None = None
            if identity.is_file() or identity.is_symlink():
                try:
                    backup_identity = read_identity(identity)
                except ValueError as exc:
                    raise DomainError(f"备份失败，身份文件无效：{exc}") from exc
                if (
                    backup_identity.database_id != self.runtime.identity.database_id
                    or backup_identity.user_id != self.runtime.identity.user_id
                    or backup_identity.device_id != self.runtime.identity.device_id
                ):
                    raise DomainError("备份失败，身份文件与当前资料库不匹配")
                backup_identity_payload = json.dumps(
                    self.runtime.identity.to_dict(), ensure_ascii=False, indent=2
                )
            asset_files: list[tuple[Path, str]] = []
            if asset_dir.is_dir():
                try:
                    asset_files = [
                        (file, f"assets/{file.relative_to(asset_dir).as_posix()}")
                        for file in iter_regular_files(asset_dir)
                    ]
                except ArchiveSecurityError as exc:
                    raise DomainError(f"备份失败，资源目录不安全：{exc}") from exc
            checksums = {"database/error_book.db": _sha256_file(snapshot)}
            if backup_identity_payload is not None:
                checksums["identity.json"] = hashlib.sha256(
                    backup_identity_payload.encode("utf-8")
                ).hexdigest()
            checksums.update(
                {archive_name: _sha256_file(file) for file, archive_name in asset_files}
            )
            manifest["checksums"] = checksums
            descriptor, archive_temp_name = tempfile.mkstemp(
                prefix=f".{dest.name}.", suffix=".tmp", dir=dest.parent
            )
            os.close(descriptor)
            archive_temp = Path(archive_temp_name)
            with zipfile.ZipFile(archive_temp, "w", compression=zipfile.ZIP_DEFLATED) as zf:
                zf.writestr("manifest.json", json.dumps(manifest, ensure_ascii=False, indent=2))
                zf.write(snapshot, arcname="database/error_book.db")
                if backup_identity_payload is not None:
                    zf.writestr("identity.json", backup_identity_payload)
                for file, archive_name in asset_files:
                    zf.write(file, arcname=archive_name)
            with archive_temp.open("r+b") as stream:
                stream.flush()
                os.fsync(stream.fileno())
            os.replace(archive_temp, dest)
            archive_temp = None
            return dest
        finally:
            if archive_temp is not None:
                archive_temp.unlink(missing_ok=True)
            shutil.rmtree(staging, ignore_errors=True)

    def restore_backup(self, zip_path: Path, target_root: Path) -> Path:
        zip_path = Path(zip_path)
        target_root = Path(target_root)
        if not zip_path.is_file():
            raise DomainError("备份文件不存在")
        target_root.mkdir(parents=True, exist_ok=True)
        tmp = Path(tempfile.mkdtemp(prefix=".restore-extract-", dir=target_root))
        final_staging = Path(tempfile.mkdtemp(prefix=".restore-final-", dir=target_root))
        previous = Path(tempfile.mkdtemp(prefix=".restore-previous-", dir=target_root))

        try:
            with zipfile.ZipFile(zip_path, "r") as zf:
                try:
                    infos = validate_zip_members(zf)
                except ArchiveSecurityError as exc:
                    raise DomainError(f"备份 ZIP 安全校验失败：{exc}") from exc
                names = {info.filename for info in infos}
                if "manifest.json" not in names or "database/error_book.db" not in names:
                    raise DomainError("无效的备份包")
                try:
                    manifest = json.loads(
                        read_zip_member_limited(
                            zf,
                            "manifest.json",
                            max_bytes=MAX_BACKUP_METADATA_BYTES,
                        ).decode("utf-8")
                    )
                except (
                    ArchiveSecurityError,
                    UnicodeDecodeError,
                    json.JSONDecodeError,
                ) as exc:
                    raise DomainError("备份 manifest.json 无效") from exc
                if (
                    not isinstance(manifest, dict)
                    or manifest.get("format") != "yancuo-local-backup"
                ):
                    raise DomainError("备份格式不匹配")
                try:
                    backup_version = int(manifest.get("version") or 0)
                    package_schema = int(manifest.get("schema_version") or 0)
                except (TypeError, ValueError) as exc:
                    raise DomainError("备份 manifest 版本字段无效") from exc
                if backup_version not in {1, 2}:
                    raise DomainError("备份版本不受支持")
                from yancuo_win.domain.identity import SCHEMA_VERSION

                if package_schema > SCHEMA_VERSION:
                    raise DomainError(
                        f"备份 schema_version={package_schema} 高于程序支持的 {SCHEMA_VERSION}，请升级软件"
                    )
                try:
                    safe_extract_zip(zf, tmp)
                except ArchiveSecurityError as exc:
                    raise DomainError(f"备份 ZIP 解压被拒绝：{exc}") from exc

            try:
                _verify_local_backup_checksums(tmp, manifest)
            except ArchiveSecurityError as exc:
                raise DomainError(f"备份 checksum 校验失败：{exc}") from exc

            db_src = tmp / "database" / "error_book.db"
            assets_src = tmp / "assets"
            identity_src = tmp / "identity.json"
            if not db_src.is_file():
                raise DomainError("备份缺少数据库文件")
            if identity_src.is_file():
                try:
                    restored_identity = read_identity(identity_src)
                except ValueError as exc:
                    raise DomainError(f"备份身份文件无效：{exc}") from exc
                manifest_database_id = manifest.get("database_id")
                if (
                    not isinstance(manifest_database_id, str)
                    or restored_identity.database_id != manifest_database_id
                ):
                    raise DomainError("备份身份文件与 manifest 的资料库不匹配")

            shutil.copy2(db_src, final_staging / "error_book.db")
            if assets_src.is_dir():
                try:
                    from yancuo_win.infrastructure.archive import copy_tree_no_symlinks

                    copy_tree_no_symlinks(assets_src, final_staging / "assets")
                except ArchiveSecurityError as exc:
                    raise DomainError(f"备份资源目录不安全：{exc}") from exc
            else:
                (final_staging / "assets" / "objects").mkdir(parents=True)
            if identity_src.is_file():
                shutil.copy2(identity_src, final_staging / "identity.json")

            # 在替换目标目录前打开 staging 数据库并执行迁移/核心表校验，
            # 这样损坏或过旧的普通 zip 也不会覆盖一个可用的数据根。
            from yancuo_win.data.db import make_engine
            from yancuo_win.data.migrate import migrate, verify_core_tables

            try:
                staged_engine = make_engine(final_staging / "error_book.db")
                try:
                    migrate(staged_engine)
                    missing = verify_core_tables(staged_engine)
                finally:
                    staged_engine.dispose()
            except DomainError:
                raise
            except Exception as exc:
                raise DomainError(f"备份数据库校验失败：{exc}") from exc
            if missing:
                raise DomainError(f"备份数据库缺少核心表：{', '.join(missing)}")

            db_dest = target_root / "error_book.db"
            assets_dest = target_root / "assets"
            identity_dest = target_root / "identity.json"
            destinations = [db_dest, assets_dest]
            if identity_src.is_file():
                destinations.append(identity_dest)
            moved_old: list[tuple[Path, Path]] = []
            moved_new: list[Path] = []
            try:
                for destination in destinations:
                    if destination.exists() or destination.is_symlink():
                        old = previous / destination.name
                        shutil.move(str(destination), str(old))
                        moved_old.append((destination, old))
                for name in ("error_book.db", "assets"):
                    source = final_staging / name
                    destination = target_root / name
                    shutil.move(str(source), str(destination))
                    moved_new.append(destination)
                identity_final = final_staging / "identity.json"
                if identity_final.is_file():
                    shutil.move(str(identity_final), str(identity_dest))
                    moved_new.append(identity_dest)
            except Exception:
                for destination in reversed(moved_new):
                    try:
                        if destination.is_dir() and not destination.is_symlink():
                            shutil.rmtree(destination)
                        else:
                            destination.unlink(missing_ok=True)
                    except OSError:
                        pass
                for destination, old in reversed(moved_old):
                    if old.exists() or old.is_symlink():
                        shutil.move(str(old), str(destination))
                raise
            else:
                shutil.rmtree(previous, ignore_errors=True)
            return target_root
        finally:
            shutil.rmtree(tmp, ignore_errors=True)
            shutil.rmtree(final_staging, ignore_errors=True)
            shutil.rmtree(previous, ignore_errors=True)

    # ---- word export ----

    def export_problems_docx(self, problem_ids: list[str], dest: Path) -> Path:
        try:
            from docx import Document
            from docx.shared import Inches
        except ImportError as exc:  # pragma: no cover
            raise DomainError("未安装 python-docx，无法导出 Word") from exc

        problems = []
        for pid in problem_ids:
            p = self.get_problem(pid)
            if p and p.status != "trashed":
                problems.append(p)
        if not problems:
            raise DomainError("没有可导出的题目")

        doc = Document()
        doc.add_heading("研错库导出", level=0)
        for idx, p in enumerate(problems, start=1):
            title = p.title or f"题目 {idx}"
            doc.add_heading(f"{idx}. {title}", level=1)
            meta = f"优先级：{p.priority}　状态：{p.status}　ID：{p.id}"
            doc.add_paragraph(meta)
            doc.add_heading("原题", level=2)
            from yancuo_win.application.question_content import load_question_content

            content_blocks = load_question_content(p.question_content_json)
            assets_by_id = {asset.id: asset for asset in (p.assets or [])}
            if content_blocks:
                for block in content_blocks:
                    kind = block.get("type")
                    if kind in {"text", "formula"}:
                        doc.add_paragraph(str(block.get("content") or ""))
                    elif kind == "table":
                        rows = block.get("rows") or []
                        column_count = max(
                            (
                                sum(
                                    max(1, int(cell.get("colspan", 1)))
                                    if isinstance(cell, dict)
                                    else 1
                                    for cell in row
                                )
                                for row in rows
                                if isinstance(row, list)
                            ),
                            default=1,
                        )
                        table = doc.add_table(rows=max(1, len(rows)), cols=max(1, column_count))
                        occupied: set[tuple[int, int]] = set()
                        for row_index, row in enumerate(rows):
                            column_index = 0
                            for raw_cell in row:
                                while (row_index, column_index) in occupied:
                                    column_index += 1
                                cell_data = raw_cell if isinstance(raw_cell, dict) else {}
                                content = (
                                    str(cell_data.get("content") or "")
                                    if cell_data
                                    else str(raw_cell or "")
                                )
                                rowspan = max(1, int(cell_data.get("rowspan", 1)))
                                colspan = max(1, int(cell_data.get("colspan", 1)))
                                end_row = min(len(rows) - 1, row_index + rowspan - 1)
                                end_col = min(column_count - 1, column_index + colspan - 1)
                                target = table.cell(row_index, column_index)
                                if end_row != row_index or end_col != column_index:
                                    target = target.merge(table.cell(end_row, end_col))
                                target.text = content
                                for used_row in range(row_index, end_row + 1):
                                    for used_col in range(column_index, end_col + 1):
                                        occupied.add((used_row, used_col))
                                column_index = end_col + 1
                    elif kind == "figure":
                        asset = assets_by_id.get(str(block.get("derived_asset_id") or ""))
                        if asset is not None:
                            path = self.store.resolve(asset.relative_path)
                            if path.is_file():
                                doc.add_picture(str(path), width=Inches(6.0))
                        caption = str(block.get("content") or "").strip()
                        if caption:
                            doc.add_paragraph(caption)
            else:
                doc.add_paragraph(p.question_markdown or "（空）")
                if p.question_latex:
                    doc.add_paragraph(f"LaTeX：{p.question_latex}")
            doc.add_heading("我的作答", level=2)
            doc.add_paragraph(p.user_answer or "（空）")
            doc.add_heading("正确答案", level=2)
            doc.add_paragraph(p.correct_answer or "（空）")
            doc.add_heading("解析", level=2)
            doc.add_paragraph(p.solution_markdown or "（空）")
            if p.error_analysis:
                doc.add_heading("错因", level=2)
                doc.add_paragraph(p.error_analysis)
            if p.notes:
                doc.add_heading("备注", level=2)
                doc.add_paragraph(p.notes)
            doc.add_paragraph("")

        dest = Path(dest)
        dest.parent.mkdir(parents=True, exist_ok=True)
        doc.save(dest)
        return dest
